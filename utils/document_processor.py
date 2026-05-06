import os
import chardet
from loguru import logger
from typing import List
from pathlib import Path


def extract_text_from_file(file_path: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT files."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext in (".txt", ".md"):
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract table text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_txt(file_path: str) -> str:
    with open(file_path, "rb") as f:
        raw = f.read()
    encoding = chardet.detect(raw).get("encoding", "utf-8") or "utf-8"
    return raw.decode(encoding, errors="replace")


def chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 50,
) -> List[str]:
    """Split text into overlapping chunks by word count."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if len(chunk.strip()) > 20:  # skip tiny fragments
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks
